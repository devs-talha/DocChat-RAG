import os
import chainlit as cl
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_openai import ChatOpenAI
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain import hub
from chainlit import make_async
from config import settings

retrieval_qa_chat_prompt = hub.pull(settings.retrieval_qa_chat_prompt)

def build_vector_store(docs: list[Document]):
    """Return a retriever built on a fresh in-memory FAISS vector store for the given docs."""
    embeddings = HuggingFaceEmbeddings(model_name=settings.embedding_model)
    vector_store = FAISS.from_documents(docs, embeddings)
    retriever = vector_store.as_retriever(search_type=settings.retriever_search_type, search_kwargs={"k": settings.retriever_k})
    return retriever


async_build_vector_store = make_async(build_vector_store)


def docs_from_filepaths(file_paths: list[str]) -> list[Document]:
    """Load text from given file paths (txt, md, csv, pdf, docx) and split into Documents."""

    import io
    splitter = RecursiveCharacterTextSplitter(chunk_size=settings.chunk_size, chunk_overlap=settings.chunk_overlap)
    documents: list[Document] = []

    for path in file_paths:
        ext = os.path.splitext(path)[1].lower()
        text_content = ""
        if ext in {".txt", ".md", ".csv"}:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text_content = f.read()
        elif ext == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    pages_text = [p.extract_text() or "" for p in pdf.pages]
                    text_content = "\n".join(pages_text)
            except ImportError:
                text_content = ""  # Skip if pdfplumber missing
        elif ext == ".docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(path)
                text_content = "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                text_content = ""
        else:
            continue  # Unsupported type

        if not text_content:
            continue

        metadata = {"source": os.path.basename(path)}
        chunks = splitter.split_text(text_content)
        documents.extend([Document(page_content=ch, metadata=metadata) for ch in chunks])
    return documents

async_docs_from_filepaths = make_async(docs_from_filepaths)

@cl.on_chat_start
async def on_chat_start():
    msg = cl.AskFileMessage(
        "👋 Hi! Upload one or more files and I will answer questions using their content.",
        accept={
            "text/plain": [".txt", ".md", ".csv"],
            "application/pdf": [".pdf"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": [".docx"]
        },
        max_size_mb=settings.max_file_size_mb,
        max_files=settings.max_files,
    )
    files = await msg.send()  # Wait for user upload
    
    msg.content = "Processing files..."
    await msg.update()
    
    if files is None:
        await cl.Message(content="No files uploaded. Please try again.").send()
        return await on_chat_start()

    # Save temp files and build docs
    file_paths = []
    for file in files:
        tmp_path = file.path
        file_paths.append(tmp_path)

    docs = await async_docs_from_filepaths(file_paths)
    retriever = await async_build_vector_store(docs)
    # Store retriever in the user session so each chat has its own knowledge base
    cl.user_session.set("retriever", retriever)

    msg.content = "I am ready! Ask me anything!"
    await msg.update()


@cl.on_message
async def on_message(msg: cl.Message):
    retriever = cl.user_session.get("retriever")
    if retriever is None:
        await cl.Message(content="❗️ No documents indexed yet. Please restart the chat and upload files.").send()
        return

    llm = ChatOpenAI(temperature=settings.openai_temperature, streaming=True, max_tokens=settings.openai_max_tokens, model=settings.openai_model)
    
    combine_docs_chain = create_stuff_documents_chain(
        llm, retrieval_qa_chat_prompt
    )

    retrieval_chain = create_retrieval_chain(retriever, combine_docs_chain)

    answer_msg = cl.Message(content="")
    await answer_msg.send()
    collected_answer = ""
    async for chunk in retrieval_chain.astream({"input": msg.content}):
        token = chunk.get("answer", "")
        if token:
            collected_answer += token
            await answer_msg.stream_token(token)
    answer_msg.content = collected_answer
    await answer_msg.update()